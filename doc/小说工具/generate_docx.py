#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
指尖仙侠详细功能说明书 - DOCX生成脚本
插入截图 + 结构化排版
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# 配置
BASE_DIR = r"C:\Users\rock\Doubao\chats\2026-08-01\new-chat"
OUTPUT_FILE = os.path.join(BASE_DIR, "01-详细功能说明书-v2.docx")

def set_chinese_font(run, font_name="微软雅黑", size=10.5):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(size)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_chinese_font(run, size=18 - level * 2)
    return heading

def add_paragraph(doc, text, bold=False, size=10.5):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    set_chinese_font(run, size=size)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        set_chinese_font(run, size=9)
    
    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_chinese_font(run, size=9)
    
    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)
    
    return table

def add_image(doc, image_path, width=6.0, caption=None):
    """添加图片"""
    if os.path.exists(image_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_path, width=Inches(width))
            
            if caption:
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_p.add_run(caption)
                cap_run.italic = True
                set_chinese_font(cap_run, size=9)
                cap_run.font.color.rgb = RGBColor(128, 128, 128)
            return True
        except Exception as e:
            add_paragraph(doc, f"[图片加载失败: {os.path.basename(image_path)} - {e}]", size=9)
            return False
    else:
        add_paragraph(doc, f"[图片不存在: {os.path.basename(image_path)}]", size=9)
        return False

def img(filename):
    """获取图片完整路径"""
    return os.path.join(BASE_DIR, filename)

def main():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ========== 封面 ==========
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_heading('指尖仙侠', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_chinese_font(run, size=28)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('详细功能说明书')
    run.bold = True
    set_chinese_font(run, size=18)
    
    doc.add_paragraph()
    
    info_lines = [
        '版本：v2.0',
        '生成日期：2026-08-01',
        '数据来源：功能手册 + 技术手册 + 界面截图 + 代码级验证',
        '文档性质：完整功能全景手册',
    ]
    
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_chinese_font(run, size=11)
    
    doc.add_page_break()
    
    # ========== 目录页 ==========
    add_heading(doc, '目录', level=1)
    
    toc_items = [
        '一、产品概览',
        '二、界面与导航',
        '三、世界观与设定',
        '四、大纲与场景脚本',
        '五、生成控制台',
        '六、章节阅读与编辑',
        '七、叙事工程',
        '八、质量防线',
        '九、素材与配置',
        '十、数据架构',
        '附录',
    ]
    
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_chinese_font(run, size=11)
    
    doc.add_page_break()
    
    # ========== 一、产品概览 ==========
    add_heading(doc, '一、产品概览', level=1)
    
    add_heading(doc, '1.1 一句话定位', level=2)
    add_paragraph(doc, 
        '指尖仙侠是一个本地化部署的 AI 长篇小说创作工具，以结构化世界观数据库为知识源，'
        '通过多 Agent 管线自动生成符合设定的小说章节，并提供完整的叙事工程与质量防线体系。')
    
    add_heading(doc, '1.2 核心能力', level=2)
    add_table(doc, 
        ['能力维度', '说明'],
        [
            ['结构化世界观', '诛仙小说 18 表结构化数据库 + pgvector 512 维向量，支持语义检索'],
            ['多 Agent 生成管线', 'ContextComposer → Writer → Auditor → Reviser，4 步流水线'],
            ['叙事工程体系', '伏笔台账 + 因果链 + 时间线 + 方向统计 + 影响体系 + 状态追踪'],
            ['质量防线', '27 维剧情审计 + 8 维文风校验 + 9 维叙事体检 + 去 AI 味检测'],
            ['自定义工坊', '自定义人物/功法/法宝，支持千人千面变种系统'],
            ['降级红线架构', '所有增强点 best-effort，失败不阻断主流程'],
        ],
        col_widths=[1.5, 4.5]
    )
    
    add_heading(doc, '1.3 技术栈', level=2)
    add_table(doc,
        ['层级', '技术'],
        [
            ['前端', 'React 19 + TypeScript + Tailwind CSS + @dnd-kit'],
            ['后端', 'Hono + TypeScript + Drizzle ORM'],
            ['数据库', 'PostgreSQL + pgvector（HNSW cosine 索引）'],
            ['向量模型', 'BAAI/bge-small-zh-v1.5，512 维'],
            ['Python 旁路', 'Embedding server（8600）+ ETL/GUI server（8610）'],
            ['包管理', 'pnpm monorepo'],
            ['部署方式', '本地化部署，支持 Windows/macOS/Linux'],
        ],
        col_widths=[1.5, 4.5]
    )
    
    add_heading(doc, '1.4 功能全景图（4 条产品线）', level=2)
    add_paragraph(doc, '产品分为四条产品线：A 创作主线、B 叙事工程、C 质量防线、D 素材与配置。')
    
    add_table(doc,
        ['A 创作主线', 'B 叙事工程', 'C 质量防线', 'D 素材与配置'],
        [
            ['仪表盘', '伏笔台账', '叙事体检', '金句库'],
            ['世界观百科', '因果链', '27维质量审计', '热点嗅探'],
            ['大纲与场景', '时间线', '文风校验', '素材知识库'],
            ['生成控制台', '方向统计', '去AI味检测', '自定义人物'],
            ['章节阅读编辑', '影响体系', '', '自定义功法'],
            ['剧情分支', '状态追踪', '', '自定义法宝'],
            ['', '', '', '系统设置'],
        ],
        col_widths=[1.5, 1.5, 1.5, 1.5]
    )
    
    add_heading(doc, '1.5 双库架构', level=2)
    add_paragraph(doc, '产品采用双数据库设计：诛仙库（novel_db）+ 创作库（novel_studio）。')
    
    add_table(doc,
        ['数据库', '用途', '表数量', '连接池'],
        [
            ['novel_db（诛仙库）', '诛仙小说结构化数据，作为世界观知识源', '18 张', 'max 5'],
            ['novel_studio（创作库）', '用户创作数据、项目、章节、伏笔等', '40+ 张', 'max 10'],
        ],
        col_widths=[1.8, 2.2, 1.0, 1.0]
    )
    
    add_paragraph(doc, '关键设计：', bold=True)
    add_paragraph(doc, '• 自定义实体使用负数 ID，与诛仙库正数 ID 共存分流')
    add_paragraph(doc, '• 诛仙库非严格只读，存在写入操作（如蒸馏结果）')
    add_paragraph(doc, '• 两库独立连接池，无数据库级外键关联')
    add_paragraph(doc, '• 应用层通过 project_id 实现项目隔离')
    
    doc.add_page_break()
    
    # ========== 二、界面与导航 ==========
    add_heading(doc, '二、界面与导航', level=1)
    
    add_heading(doc, '2.1 侧边栏 / 信息架构', level=2)
    
    add_image(doc, img('sidebar.png'), width=2.0, caption='图 2-1：侧边栏导航')
    
    add_paragraph(doc, '代码确认事实：', bold=True)
    add_paragraph(doc, '• 共 16 个菜单项，完全平铺，无分组、无二级菜单')
    add_paragraph(doc, '• 来源：App.tsx L42-59 navItems 数组')
    add_paragraph(doc, '• 排序是硬编码书写顺序，无动态排序逻辑')
    
    add_paragraph(doc, '16 个菜单项：', bold=True)
    add_table(doc,
        ['序号', '菜单项', '产品线', '说明'],
        [
            ['1', '仪表盘', 'A', '首页/驾驶舱'],
            ['2', '众生百态', 'D', '自定义人物管理'],
            ['3', '铸器天工', 'D', '自定义法宝管理'],
            ['4', '道法自然', 'D', '自定义功法管理'],
            ['5', '世界观', 'A', '诛仙库世界观百科'],
            ['6', '大纲', 'A', '卷大纲 + 场景脚本'],
            ['7', '生成', 'A', '生成控制台'],
            ['8', '章节', 'A', '章节阅读与编辑'],
            ['9', '叙事线索', 'B', '伏笔台账 + 因果链'],
            ['10', '时间线', 'B', '剧情时间线 + 人物状态'],
            ['11', '金句库', 'D', '金句 + 名场面'],
            ['12', '叙事体检', 'C', '全书规则体检'],
            ['13', '方向统计', 'B', '剧情方向分布'],
            ['14', '热点嗅探', 'D', '榜单抓取 + 灵感'],
            ['15', '素材知识库', 'D', 'RAG 知识库 + ETL'],
            ['16', '设置', 'D', '系统设置'],
        ],
        col_widths=[0.5, 1.2, 0.8, 3.5]
    )
    
    add_paragraph(doc, '我的理解：', bold=True)
    add_paragraph(doc, '当前信息架构问题较大：16 项平铺无分组，命名三种风格混用（文艺仙侠风/直白功能风/混合风），'
                       '顺序不符合创作工作流（三个自定义模块放在世界观前面，逻辑反了）。'
                       '质量防线模块散落各处，缺少全局搜索（Cmd+K）、最近访问等高效导航方式。')
    
    add_heading(doc, '2.2 顶部栏', level=2)
    add_paragraph(doc, '顶部栏元素：')
    add_paragraph(doc, '• 左侧：Logo + 产品名')
    add_paragraph(doc, '• 中部：双库状态指示（诛仙库/创作库 连接状态）')
    add_paragraph(doc, '• 右侧：卷筛选下拉框（默认"全书"）')
    add_paragraph(doc, '"双库状态"是连接状态文字（连通/断开），不是进度条。')
    
    add_heading(doc, '2.3 仪表盘 / 首页', level=2)
    
    add_image(doc, img('dashboard.png'), width=6.0, caption='图 2-2：仪表盘首页')
    
    add_paragraph(doc, '代码确认事实：', bold=True)
    add_paragraph(doc, '• 实际 6 个卡片：双库状态 / 创作热力图 / 伏笔回收进度 / 方向均衡度 / 我的项目 / 最近生成任务')
    add_paragraph(doc, '• 创作热力图是相对分级（按当前最大值比例映射），不是固定阈值')
    add_paragraph(doc, '• 热力图只有 hover tooltip，不能点击下钻')
    add_paragraph(doc, '• "我的项目"可点击切换 project_id')
    
    add_paragraph(doc, '6 个卡片详情：', bold=True)
    add_table(doc,
        ['卡片', '内容', '数据来源'],
        [
            ['双库状态', '诛仙库/创作库连接状态', '数据库连接池'],
            ['创作热力图', '类 GitHub 贡献图，按日字数渲染', '创作统计表'],
            ['伏笔回收进度', '已回收/总数 + 超期数量', '伏笔台账'],
            ['方向均衡度', '均衡度评分 + 占比最高大类', '方向统计'],
            ['我的项目', '项目列表，点击切换', '项目表'],
            ['最近生成任务', '近期任务状态/耗时/tokens', '生成队列表'],
        ],
        col_widths=[1.5, 2.5, 2.0]
    )
    
    add_paragraph(doc, '我的理解：', bold=True)
    add_paragraph(doc, '仪表盘定位有问题：是"数据展示板"而不是"驾驶舱+行动入口"。'
                       '最大面积给了创作热力图，但热力图价值最低。'
                       '缺少全书概览关键指标、待办提醒聚合、快速操作入口。')
    
    doc.add_page_break()
    
    # ========== 三、世界观与设定 ==========
    add_heading(doc, '三、世界观与设定', level=1)
    
    add_heading(doc, '3.1 世界观百科', level=2)
    
    add_image(doc, img('worldview1.png'), width=6.0, caption='图 3-1：世界观百科 - 人物列表')
    
    add_paragraph(doc, '代码确认事实：', bold=True)
    add_paragraph(doc, '• 前端入口：/world（WorldBrowser.tsx，1541 行，全内联）')
    add_paragraph(doc, '• 后端接口：/world')
    add_paragraph(doc, '• 顶部：BookSwitcher（书籍切换，默认 bookId=1）+ 数据总览 chips')
    
    add_heading(doc, '3.1.1 8 类实体', level=3)
    add_table(doc,
        ['序号', '实体类型', '说明', 'CRUD'],
        [
            ['1', '人物', '角色设定', '✅ 新增/编辑/软删除'],
            ['2', '门派', '宗门设定', '✅'],
            ['3', '地点', '地理设定', '✅'],
            ['4', '功法', '功法设定', '✅'],
            ['5', '法宝', '法宝设定', '✅'],
            ['6', '妖兽', '妖兽设定', '✅'],
            ['7', '丹药灵材', '丹药和灵材', '✅'],
            ['8', '日常信物', '日常物品和信物', '✅'],
        ],
        col_widths=[0.5, 1.2, 2.5, 1.8]
    )
    add_paragraph(doc, '列表展示形式：卡片网格')
    
    add_heading(doc, '3.1.2 全局搜索（已知 bug）', level=3)
    add_paragraph(doc, '⚠️ 已知 bug：', bold=True)
    add_paragraph(doc, '• 后端 world.ts L414 返回对象格式')
    add_paragraph(doc, '• 前端 WorldBrowser.tsx L1451 用 Array.isArray() 判断')
    add_paragraph(doc, '• 判断失败 → 渲染为空 → 搜索完全不可用')
    add_paragraph(doc, '• 影响：925 篇功法、数百个人物，没有搜索根本找不到')
    
    add_heading(doc, '3.1.3 人物详情', level=3)
    
    add_image(doc, img('worldview2.png'), width=6.0, caption='图 3-2：世界观百科 - 人物详情')
    
    add_paragraph(doc, '人物详情 3 个 Tab：', bold=True)
    add_table(doc,
        ['Tab', '内容', '说明'],
        [
            ['设定', '基本设定', '姓名、性别、境界、身份、性格、外貌、背景...'],
            ['关系', '人物关系', '关系类型、对方人物、关系描述、互动次数'],
            ['深度蒸馏', 'AI 提炼的人物内核', '心智模型、决策启发式、人生阶段'],
        ],
        col_widths=[1.0, 1.5, 3.5]
    )
    
    add_paragraph(doc, '关于"深度蒸馏"：', bold=True)
    add_paragraph(doc, '这是产品化命名，实质是 LLM 蒸馏的写作画像。'
                       '当前展示方式：直接 JSON.stringify 原始输出，无可视化。'
                       '用户看到一大段 JSON 文本，体验极差。'
                       '建议改名："道心"或"人物内核"（更仙侠、更通俗）。')
    
    add_paragraph(doc, '人物关系数据模型：', bold=True)
    add_paragraph(doc, '• 诛仙库关系：relType 自由文本 + interactCount 互动次数')
    add_paragraph(doc, '• 无亲密度字段（自定义人物关系才有 relLevel 1-5 数值）')
    add_paragraph(doc, '• 只有列表，无关系图谱')
    
    add_heading(doc, '3.1.4 功法详情', level=3)
    add_paragraph(doc, '功法详情 2 个 Tab：设定 / 蒸馏')
    add_paragraph(doc, '诛仙库功法招式用独立表 technique_move，结构化存储。')
    
    add_heading(doc, '3.1.5 文风引擎', level=3)
    add_paragraph(doc, '文风引擎是世界观内的子模块，共 8 个模块：')
    add_table(doc,
        ['模块', '说明'],
        [
            ['心智模型', '叙事视角和思维方式'],
            ['决策启发式', '角色决策模式'],
            ['描写比例', '各类描写的占比（环境/心理/动作/对话）'],
            ['核心意象', '常用意象和象征'],
            ['禁用词', '禁止使用的词汇'],
            ['视角规则', '叙事视角规则（POV 规则、视角切换规则）'],
            ['反模式', '避免的写作模式'],
            ['场景文风映射', '不同场景的文风预设'],
        ],
        col_widths=[1.5, 4.5]
    )
    add_paragraph(doc, '⚠️ 文风引擎是书籍级（bookId），但生成时硬编码 bookId=1。')
    
    add_heading(doc, '3.2 自定义人物（众生百态）', level=2)
    add_paragraph(doc, '• 独立模块，不在 WorldBrowser 里')
    add_paragraph(doc, '• 自定义人物使用负数 ID，与诛仙库正数 ID 共存分流')
    add_paragraph(doc, '• 人物关系有 relLevel（1-5 数值），诛仙库没有')
    add_paragraph(doc, '• 先天禀赋：4 大类 166 条正向 + 23 条小缺陷')
    
    add_heading(doc, '3.3 自定义功法坊（道法自然）', level=2)
    
    add_paragraph(doc, '代码确认事实：', bold=True)
    add_paragraph(doc, '• 前端入口：CustomTechniqueForge.tsx，1130 行，全内联')
    add_paragraph(doc, '• 列表是卡片网格，无筛选排序，无搜索框')
    
    add_heading(doc, '3.3.1 功法详情', level=3)
    add_paragraph(doc, '⚠️ 重要纠偏：详情弹窗根本没有 Tab！', bold=True)
    add_paragraph(doc, '详情是单个可滚动 Dialog，由多个 DetailBlock 顺序铺陈：')
    
    add_table(doc,
        ['序号', '字段块', '说明'],
        [
            ['1', '道则构型', '主道则 + 辅修道则'],
            ['2', '本源运用方向', '核心方向 + 风格类型'],
            ['3', '适配门槛', '修炼条件'],
            ['4', '分道境神通', '各境界的神通'],
            ['5', '典型运用技巧', '常用技巧'],
            ['6', '反噬代价', '修炼反噬'],
            ['7', '演化方向', '功法演化路径'],
            ['8', '身体印记', '修炼后身体变化'],
            ['9', '功法详解', '纯文本段落'],
        ],
        col_widths=[0.5, 1.5, 4.0]
    )
    
    add_heading(doc, '3.3.2 千人千面·变种系统', level=3)
    add_paragraph(doc, '核心亮点功能：', bold=True)
    add_paragraph(doc, '• 基于四因子确定性推导：主道则 × 种族 × 性格 × 出身')
    add_paragraph(doc, '• 变种稀有度：普通 / 显著 / 稀有异变')
    
    add_paragraph(doc, '数据模型亮点：', bold=True)
    add_paragraph(doc, '变种差异存储（offsets），只存相对原功法的偏移量。'
                       '不存完整副本，节省空间。原功法修改后，变种自动更新。'
                       '这个设计非常聪明。')
    
    add_heading(doc, '3.3.3 创建向导', level=3)
    add_paragraph(doc, '三步创建向导：道则构型 → 行功根骨 → 衍化配置')
    add_paragraph(doc, '有"一键全随机"按钮，按兼容矩阵随机配比（高/中/冲约 60:35:5）')
    add_paragraph(doc, '道则选择有对冲警告（冲突标红）')
    
    doc.add_page_break()
    
    # ========== 四、大纲与场景脚本 ==========
    add_heading(doc, '四、大纲与场景脚本', level=1)
    
    add_heading(doc, '4.1 概述', level=2)
    add_paragraph(doc, '两个 Tab：卷大纲 / 场景脚本，共享同一份 outlines 数据。')
    
    add_heading(doc, '4.2 卷大纲', level=2)
    
    add_image(doc, img('outline1.png'), width=6.0, caption='图 4-1：卷大纲')
    
    add_paragraph(doc, '卷大纲字段：标题、概要、关键事件、人物弧线等')
    add_paragraph(doc, '章节计划 CRUD 在大纲模块内完成')
    
    add_heading(doc, '4.3 场景脚本', level=2)
    
    add_image(doc, img('scene_script1.png'), width=6.0, caption='图 4-2：场景脚本三栏布局')
    
    add_paragraph(doc, '三栏布局：', bold=True)
    add_table(doc,
        ['区域', '宽度', '内容'],
        [
            ['左侧素材池', '220px', '8 类素材（含伏笔）'],
            ['中间场景画布', '自适应', '场景节点列表'],
            ['右侧校验侧栏', '300px', '一致性校验结果（条件显示）'],
        ],
        col_widths=[1.5, 1.5, 3.0]
    )
    
    add_heading(doc, '4.3.1 一致性校验', level=3)
    add_paragraph(doc, '⚠️ 纠偏：一致性校验是 8 维（不是 7 维）：', bold=True)
    add_table(doc,
        ['序号', '维度', '说明'],
        [
            ['1', '时间线', '时间线一致性'],
            ['2', '地点', '地点一致性'],
            ['3', '结构节奏', '结构节奏合理性'],
            ['4', '人物出场', '人物出场合理性'],
            ['5', '战斗', '战斗合理性（占位）'],
            ['6', '故事引擎相关性', '故事引擎关联度'],
            ['7', '场景有效性', '场景有效性'],
            ['8', '节奏健康度', '节奏健康度'],
        ],
        col_widths=[0.5, 1.5, 4.0]
    )
    add_paragraph(doc, '校验结果右侧栏三色分级：error 红 / warning 琥珀 / info 灰')
    add_paragraph(doc, '点击可滚动定位到对应节点')
    
    doc.add_page_break()
    
    # ========== 五、生成控制台 ==========
    add_heading(doc, '五、生成控制台', level=1)
    
    add_image(doc, img('generation.png'), width=6.0, caption='图 5-1：生成控制台')
    
    add_paragraph(doc, '三栏布局比例 3:6:3')
    add_paragraph(doc, '左：章节列表 / 中：预览区 / 右：参数+队列+日志')
    
    add_heading(doc, '5.1 生成参数', level=2)
    add_paragraph(doc, 'UI 上暴露 4 个参数：')
    add_table(doc,
        ['参数', '说明', '问题'],
        [
            ['目标字数', '目标章节字数', '—'],
            ['温度', '生成随机性', '术语太技术化，无通俗解释'],
            ['自动修订', '生成后是否自动走 Reviser', '—'],
            ['文风档位', '文风预设', '硬编码 5 个选项，不可自定义'],
        ],
        col_widths=[1.2, 2.0, 2.8]
    )
    
    add_paragraph(doc, '⚠️ 重要发现：POV / 方向 / 出场人物 / 必含伏笔等参数从 chapter_plan 读取，不在控制台面板暴露。')
    
    add_heading(doc, '5.2 "生成日志"名不副实', level=2)
    add_paragraph(doc, '⚠️ 重大纠偏：右栏标着"生成日志"的卡片，实际展示的是章节计划，不是真正的步骤日志。命名误导用户。')
    
    add_heading(doc, '5.3 多 Agent 生成管线', level=2)
    add_table(doc,
        ['步骤', '名称', 'Agent/模块'],
        [
            ['1', '加载章节计划 + 奇点配额校验', '🔒 规则'],
            ['2', '构建上下文 + 裁剪 token 预算', 'context-builder'],
            ['3', 'Writer 写作', 'WriterAgent 🤖'],
            ['3.5', '本地质量预校验', 'quality-gate 🔒'],
            ['3.5.1', '先行修订', 'ReviserAgent 🤖'],
            ['4', 'Auditor 27 维审计', 'AuditorAgent 🤖'],
            ['4.1', '回炉循环（最多2轮）', 'ReviserAgent 🤖'],
            ['5.4', '精修压缩', 'condense 🤖'],
            ['6', '保存正文版本', '—'],
            ['7', '章节计划置 generated', '—'],
            ['7.5', '状态抽取', 'StateExtractorAgent 🤖'],
            ['7.55', '伏笔自动流转', 'autoUpdate 🔒'],
            ['7.6', '自动生成分支', 'BranchGeneratorAgent 🤖'],
            ['7.7', '金句提取', 'extractQuotes 🤖'],
            ['8', '完成任务', '—'],
        ],
        col_widths=[0.8, 2.5, 2.7]
    )
    
    add_heading(doc, '5.4 已知 bug', level=2)
    add_paragraph(doc, '⚠️ 单章参数丢弃 bug：', bold=True)
    add_paragraph(doc, '单章 /start 的 schema 丢弃了 targetWords/temperature/autoRevise')
    add_paragraph(doc, '来源：generation.ts L16-28')
    add_paragraph(doc, '影响：单章生成时这三个参数不生效，用户调了参数但不生效，会以为是 AI 不听话，对产品失去信任。')
    
    doc.add_page_break()
    
    # ========== 六、章节阅读与编辑 ==========
    add_heading(doc, '六、章节阅读与编辑', level=1)
    
    add_image(doc, img('chapter1.png'), width=6.0, caption='图 6-1：章节阅读与编辑')
    
    add_paragraph(doc, '三栏比例 3:6:3')
    add_paragraph(doc, '左：章节目录 / 中：正文+模式切换 / 右：剧情分支')
    
    add_heading(doc, '6.1 五模式切换', level=2)
    add_paragraph(doc, '⚠️ 纠偏：是五模式切换（不是几个 Tab），在中栏顶部：', bold=True)
    add_table(doc,
        ['模式', '说明'],
        [
            ['阅读', '阅读模式，纯文本展示'],
            ['编辑', '编辑模式，纯文本 Textarea'],
            ['版本', '版本对比，diffLines 差异对比'],
            ['修订', '对话式 AI 修订'],
            ['文风校验', '文风校验 + 去 AI 味检测'],
        ],
        col_widths=[1.2, 4.8]
    )
    
    add_heading(doc, '6.2 对话式 AI 修订', level=2)
    add_paragraph(doc, '• 选中文字 → 输入指令 → 生成 diff 预览 → 存为新版本')
    add_paragraph(doc, '• 视角重写：按指定视角重写片段，存于 perspective_versions，不覆盖原文')
    
    add_heading(doc, '6.3 文风校验', level=2)
    add_paragraph(doc, '• 8 维文风校验 + 8 类"AI 味"检测')
    add_paragraph(doc, '• 问题按 critical/major/minor 分级')
    add_paragraph(doc, '• 点击问题能定位正文（高亮）')
    add_paragraph(doc, '• 一键修正：取所有 critical/major 问题合成修订指令')
    
    add_heading(doc, '6.4 剧情分支', level=2)
    add_paragraph(doc, '右栏宽度问题：占比 3，确实偏窄，分支选项描述经常换行。')
    add_paragraph(doc, '分支选项卡片：标签 / 描述 / 影响预览（多维度） / 伏笔计数')
    add_paragraph(doc, '选错覆盖式重选，带 409 冲突保护（防并发覆盖）')
    
    add_heading(doc, '6.5 27 维质量审计', level=2)
    add_paragraph(doc, '⚠️ 重要发现：27 维质量审计只在生成管线内自动跑，章节页不能手动触发。')
    
    add_heading(doc, '6.6 版本管理', level=2)
    add_paragraph(doc, '多行版本链：version / parentVersionId / isCurrent')
    add_paragraph(doc, '存全量快照（非 diff），支持分支版本')
    
    doc.add_page_break()
    
    # ========== 七、叙事工程 ==========
    add_heading(doc, '七、叙事工程', level=1)
    
    add_heading(doc, '7.1 叙事线索（伏笔台账 + 因果链）', level=2)
    
    add_image(doc, img('foreshadow1.png'), width=6.0, caption='图 7-1：伏笔台账')
    
    add_paragraph(doc, '两个独立实体/两张表：foreshadow_thread（伏笔）+ causal_chain（因果链）')
    add_paragraph(doc, '不是同一东西的两视图，概念相近易混（产品设计问题）')
    
    add_heading(doc, '7.1.1 伏笔台账', level=3)
    add_paragraph(doc, '顶部 8 个统计数字：总数 / 待埋 / 已埋 / 已收 / 废弃 / 超期 / 分支衍生 / 待回填')
    
    add_paragraph(doc, '回填埋设（核心亮点功能）：', bold=True)
    add_table(doc,
        ['模式', '说明', '技术实现'],
        [
            ['anchor', '标记式，写入 must_have_events', '简单标记'],
            ['revise', '修订式，LCS diff 红绿差异预览 → 确认保存为新版本', '真正修改正文'],
        ],
        col_widths=[1.0, 3.0, 2.0]
    )
    add_paragraph(doc, '我的理解：这是最惊艳的功能之一，但入口藏在卡片右下角的小图标按钮里，用户根本发现不了。功能做了但用户找不到 = 白做。')
    
    add_heading(doc, '7.1.2 因果链', level=3)
    add_paragraph(doc, '• 独立页面，列表 + 顶部统计')
    add_paragraph(doc, '• 5 个状态：planted / foreshadowed / triggered / resolved / expired')
    add_paragraph(doc, '• 无图谱，列表形式')
    
    add_heading(doc, '7.2 时间线', level=2)
    add_paragraph(doc, '两个视图：剧情时间线 / 人物状态快照')
    
    add_heading(doc, '7.2.1 剧情时间线', level=3)
    add_paragraph(doc, '⚠️ 重要发现：时间格式是字符串叙事时间（如"第一天""数日后"），非可计算数值。无法做时间冲突检测。')
    
    add_heading(doc, '7.2.2 人物状态快照', level=3)
    add_paragraph(doc, '⚠️ 纠偏：状态维度是 5 维（不是只有境界）：', bold=True)
    add_table(doc,
        ['维度', '说明'],
        [
            ['位置', '当前位置'],
            ['境界', '修为境界'],
            ['伤势', '受伤状态'],
            ['心理', '心理状态'],
            ['持有物', '持有物品'],
        ],
        col_widths=[1.2, 4.8]
    )
    
    add_heading(doc, '7.3 方向统计', level=2)
    add_paragraph(doc, '顶部 3 个数字：方向均衡度 / 已选定分支 / 未分类方向')
    add_paragraph(doc, '⚠️ 点击方向不能下钻到章节')
    add_paragraph(doc, '叙事方向是 10 大类（不是 9 类），其中 3 类默认禁用')
    
    doc.add_page_break()
    
    # ========== 八、质量防线 ==========
    add_heading(doc, '八、质量防线', level=1)
    
    add_heading(doc, '8.1 叙事体检', level=2)
    
    add_image(doc, img('audit1.png'), width=6.0, caption='图 8-1：叙事体检')
    
    add_paragraph(doc, '⚠️ 重大纠偏：9 个维度的名字之前全猜错了！', bold=True)
    
    add_paragraph(doc, '真实 9 维：', bold=True)
    add_table(doc,
        ['序号', '维度', '说明'],
        [
            ['1', '目录连续性', '目录连续性检查'],
            ['2', '缓冲比', '节奏缓冲比例'],
            ['3', '伏笔生命周期', '伏笔生命周期健康度'],
            ['4', '角色状态链', '角色状态连续性'],
            ['5', '时代与实体', '时代与实体一致性'],
            ['6', '待决议事项', '待解决的问题'],
            ['7', '方向均衡度', '剧情方向均衡度'],
            ['8', '影响健康度', '影响体系健康度'],
            ['9', '因果链健康度', '因果链健康度'],
        ],
        col_widths=[0.5, 1.5, 4.0]
    )
    
    add_paragraph(doc, '前端文案问题：前端写"6 维度"是过期文案，实际是 9 维度。')
    
    add_heading(doc, '8.2 三个质检模块的关系', level=2)
    add_table(doc,
        ['模块', '范围', '触发方式', '维度'],
        [
            ['叙事体检', '全书', '手动触发', '9 维规则体检'],
            ['27 维质量审计', '单章', '生成时自动', '27 维 AI 审计'],
            ['文风校验', '单章', '章节页手动', '8 维文风 + 8 类 AI 味'],
        ],
        col_widths=[1.2, 0.8, 1.2, 2.8]
    )
    
    doc.add_page_break()
    
    # ========== 九、素材与配置 ==========
    add_heading(doc, '九、素材与配置', level=1)
    
    add_heading(doc, '9.1 金句库', level=2)
    add_paragraph(doc, '• 顶部 3 个数字：金句总数 / 已收藏 / 涉及人物')
    add_paragraph(doc, '• 按说话人分组')
    add_paragraph(doc, '• ⚠️ 没有搜索框')
    add_paragraph(doc, '• 人物感知注入：生成时按 POV/出场人物召回收藏金句注入')
    
    add_heading(doc, '9.2 热点嗅探', level=2)
    add_paragraph(doc, '顶部 3 Tab：抓取榜单 / 榜单书目 / 灵感入库')
    add_paragraph(doc, '⚠️ 纠偏：实际 5 个榜单源（不是 4 个）：', bold=True)
    add_paragraph(doc, '番茄男频玄幻 / 纵横月票 / 纵横人气 / 晋江月榜 / 起点月票')
    
    add_paragraph(doc, '5 型灵感：奇遇 / 伏笔 / 高光 / 任务 / 趋势')
    add_paragraph(doc, '⚠️ 推送是逐条（per-item），非批量；且为全局共享（project_id=NULL）')
    
    add_heading(doc, '9.3 素材知识库', level=2)
    add_paragraph(doc, '⚠️ 重大纠偏：三个 Tab 名字之前全猜错了！', bold=True)
    add_paragraph(doc, '真实 Tab：文风预设 / 领域知识 / 蒸馏任务')
    
    add_paragraph(doc, '技术架构：', bold=True)
    add_paragraph(doc, '• Python 旁路服务：Embedding server（8600）+ ETL server（8610）')
    add_paragraph(doc, '• 向量数据库：PostgreSQL + pgvector（HNSW cosine 索引）')
    add_paragraph(doc, '• Embedding 模型：BAAI/bge-small-zh-v1.5，512 维')
    add_paragraph(doc, '• ⚠️ 检索策略：纯向量检索（最低分 0.35），无混合检索/重排序')
    
    doc.add_page_break()
    
    # ========== 十、数据架构 ==========
    add_heading(doc, '十、数据架构', level=1)
    
    add_heading(doc, '10.1 双库架构', level=2)
    add_table(doc,
        ['数据库', '用途', '表数量', '连接池'],
        [
            ['novel_db（诛仙库）', '诛仙小说结构化数据', '18 张', 'max 5'],
            ['novel_studio（创作库）', '用户创作数据', '40+ 张', 'max 10'],
        ],
        col_widths=[1.8, 2.2, 1.0, 1.0]
    )
    
    add_heading(doc, '10.2 JSON 字段统计', level=2)
    add_paragraph(doc, '• 创作库约 100 个 jsonb 字段，诛仙库 17 个')
    add_paragraph(doc, '• 仅素材清洗库表建了 GIN 索引，其余 jsonb 多无索引')
    add_paragraph(doc, '• 性能隐患：无 GIN 索引的 jsonb 过滤、全书全量重算、伏笔/状态全表扫描')
    
    add_heading(doc, '10.3 版本管理', level=2)
    add_paragraph(doc, '• 章节：多行版本链（version/parentVersionId/isCurrent）')
    add_paragraph(doc, '• 其他实体：upsert + version+1 乐观锁')
    add_paragraph(doc, '• 各自实现，不统一')
    add_paragraph(doc, '• 版本存全量快照（非 diff）')
    
    add_heading(doc, '10.4 队列实现', level=2)
    add_paragraph(doc, '• DB 即队列 + 轮询')
    add_paragraph(doc, '• 无独立消息队列')
    add_paragraph(doc, '• 指数退避重试（30s/120s/480s，最多 3 次）')
    add_paragraph(doc, '• 默认并发 1')
    
    add_heading(doc, '10.5 降级红线', level=2)
    add_paragraph(doc, '架构原则：所有增强点用 best-effort try/catch 独立包裹，失败不阻断主流程。')
    add_paragraph(doc, '降级场景：向量服务宕机 → 向量召回降级为空；LLM 调用失败 → 队列层重试')
    add_paragraph(doc, '⚠️ 降级多为静默降级（仅 console.warn），用户不易察觉（体验隐患）')
    
    doc.add_page_break()
    
    # ========== 附录 ==========
    add_heading(doc, '附录', level=1)
    
    add_heading(doc, '附录 A：已知问题与限制', level=2)
    add_table(doc,
        ['问题', '影响', '优先级'],
        [
            ['世界观搜索有 bug（对象 vs 数组）', '搜索完全不可用', 'P0'],
            ['单章生成丢弃 3 个参数', '参数不生效，失去信任', 'P0'],
            ['"生成日志"名不副实', '命名误导', 'P0'],
            ['叙事体检前端"6 维度"过期文案', '文案错误', 'P1'],
            ['约 100 个 jsonb 字段，大部分无 GIN 索引', '性能隐患', 'P1'],
            ['降级多为静默降级', '用户不易察觉', 'P1'],
            ['27 维质量审计只能生成时自动跑', '可控性差', 'P1'],
            ['版本管理不统一', '架构债务', 'P2'],
            ['时间用字符串存储', '无法做冲突检测', 'P2'],
            ['纯向量检索，无混合检索', '检索质量有限', 'P2'],
        ],
        col_widths=[2.5, 2.0, 0.8]
    )
    
    add_heading(doc, '附录 B：正面发现（比预期好的）', level=2)
    add_paragraph(doc, '1. 功法变种差异存储（offsets）——设计很聪明')
    add_paragraph(doc, '2. 场景关联实体批量 IN 查询——无 N+1，性能设计到位')
    add_paragraph(doc, '3. 队列指数退避重试——容错设计不错')
    add_paragraph(doc, '4. 文风校验点击定位正文——交互比预期好')
    add_paragraph(doc, '5. 状态快照 5 维——数据模型比预期完善')
    add_paragraph(doc, '6. 金句有出处信息——比预期好')
    add_paragraph(doc, '7. 降级红线 best-effort try/catch——架构设计很稳')
    add_paragraph(doc, '8. 大纲一致性校验 8 维三色分级点击定位——功能比预期丰富')
    add_paragraph(doc, '9. 伏笔回填双模式（anchor+revise）——设计很细')
    add_paragraph(doc, '10. 对话式 AI 修订 + 视角重写——章节编辑功能比预期丰富')
    
    # 保存
    doc.save(OUTPUT_FILE)
    print(f"DOCX 已生成: {OUTPUT_FILE}")
    print(f"文件大小: {os.path.getsize(OUTPUT_FILE)} bytes")

if __name__ == '__main__':
    main()
